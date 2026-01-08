import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path
from typing import Optional
import subprocess
import threading
import os

from pypdf import PdfReader, PdfWriter
from pypdf.constants import UserAccessPermissions
from PIL import ImageTk
import pypdfium2 as pdfium

from docx import Document              # Word 出力
import pdfplumber                      # PDF からテキスト＆表抽出
from openpyxl import Workbook          # Excel 出力
from openpyxl.styles import Border, Side, Alignment  # 罫線 & 文字配置
from openpyxl.utils import get_column_letter

from page_views import PageSelectView, PageThumbnailView
from utils import find_gs, split_dnd_paths, open_folder

# ドラッグ&ドロップ対応（tkinterdnd2 があれば使う）
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    BaseTk = TkinterDnD.Tk
    DND_AVAILABLE = True
except Exception:
    BaseTk = tk.Tk
    DND_AVAILABLE = False


class PDFToolApp(BaseTk):
    INVALID_FILENAME_CHARS = '\\/:*?\"<>|'

    def __init__(self):
        super().__init__()

        self.title("PDF Utility")
        self.geometry("950x900")
        self.resizable(True, True)

        style = ttk.Style(self)

        # PDFファイルパスのリスト（結合タブで使用）
        self.pdf_paths: list[Path] = []

        # ステータスバー用
        self.status = tk.StringVar(value="準備完了")

        # 出力フォルダ（共通用）
        self.output_dir_var = tk.StringVar(value="")

        # 結合タブのプレビュー画像保持
        self.merge_preview_image = None

        # Ghostscript パス
        self.gs_path: Optional[str] = None

        # PDF基本情報表示用
        self.info_name = tk.StringVar(value="---")
        self.info_pages = tk.StringVar(value="---")
        self.info_size = tk.StringVar(value="---")
        self.info_path = tk.StringVar(value="---")

        # 上書き一括設定
        self.overwrite_all = tk.BooleanVar(value=False)

        self.widgets()

        # いったんレイアウトを確定させてから、そのサイズを最小サイズにする
        self.update_idletasks()
        self.minsize(self.winfo_width(), self.winfo_height())

        # 初期情報パネル
        self.update_pdf_info(None)

    # ===== 共通ヘルパー：プレースホルダ & ファイル名 =====
    def init_placeholder(self, entry: tk.Entry, placeholder_text: str):
        """
        Entryに薄いグレーのプレースホルダを設定する。
        entry._placeholder / entry._has_placeholder を内部フラグとして使う。
        """
        entry._placeholder = placeholder_text
        entry._has_placeholder = False

        def on_focus_in(e):
            if getattr(entry, "_has_placeholder", False):
                entry.delete(0, "end")
                entry.config(foreground="black")
                entry._has_placeholder = False

        def on_focus_out(e):
            if not entry.get():
                self.set_placeholder(entry, entry._placeholder)

        entry.bind("<FocusIn>", on_focus_in)
        entry.bind("<FocusOut>", on_focus_out)

        # 初期表示
        self.set_placeholder(entry, placeholder_text)

    def set_placeholder(self, entry: tk.Entry, placeholder_text: str):
        """
        すでにユーザーが入力していなければ、プレースホルダを再設定する。
        """
        entry._placeholder = placeholder_text
        if getattr(entry, "_has_placeholder", False) or not entry.get():
            entry.delete(0, "end")
            entry.insert(0, placeholder_text)
            entry.config(foreground="gray")
            entry._has_placeholder = True

    def get_entry_text(self, entry: tk.Entry) -> str:
        """
        Entryから、プレースホルダを除いた実際の入力文字列を取得する。
        """
        text = entry.get().strip()
        if getattr(entry, "_has_placeholder", False):
            return ""
        return text

    # ----- タブ別のデフォルト名 -----
    def get_merge_default_name(self) -> str:
        if self.pdf_paths:
            if len(self.pdf_paths) == 1:
                return f"{self.pdf_paths[0].stem}_結合済み.pdf"
            else:
                return f"{self.pdf_paths[0].stem}_ほか{len(self.pdf_paths)-1}件_結合済み.pdf"
        return "空欄:'元ファイル名'_ほかN件_結合済み.pdf"

    def update_merge_output_placeholder(self):
        if hasattr(self, "merge_output_entry"):
            suggestion = self.get_merge_default_name()
            self.set_placeholder(self.merge_output_entry, suggestion)

    def get_split_default_name(self, mode: str, src: Path) -> str:
        suffix = "_抽出済み" if mode == "keep" else "_削除済み"
        return f"{src.stem}{suffix}.pdf"

    def update_split_output_placeholder(self):
        if not getattr(self, "split_src_path", None):
            placeholder = "空欄:'元ファイル名'_抽出済み.pdf"
        else:
            placeholder = self.get_split_default_name(self.page_edit_mode.get(), self.split_src_path)
        if hasattr(self, "split_output_entry"):
            self.set_placeholder(self.split_output_entry, placeholder)

    def get_reorder_default_name(self, in_path: Path) -> str:
        return f"{in_path.stem}_並び替え済み.pdf"

    def update_reorder_output_placeholder(self):
        if not getattr(self, "reorder_pdf_path", None):
            placeholder = "空欄:'元ファイル名'_並び替え済み.pdf"
        else:
            placeholder = self.get_reorder_default_name(Path(self.reorder_pdf_path))
        if hasattr(self, "reorder_output_entry"):
            self.set_placeholder(self.reorder_output_entry, placeholder)

    # === 圧縮タブ用 デフォルト名 ===
    def get_compress_default_name(self, src: Path) -> str:
        return f"{src.stem}_圧縮済み.pdf"

    def get_compress_default_suffix(self) -> str:
        return "空欄:'元ファイル名'_圧縮済み.pdf"

    def update_compress_suffix_placeholder(self):
        if not hasattr(self, "compress_suffix_entry"):
            return

        files = getattr(self, "compress_files", [])
        if files:
            example = self.get_compress_default_name(files[0])
            placeholder = example
        else:
            placeholder = self.get_compress_default_suffix()

        self.set_placeholder(self.compress_suffix_entry, placeholder)

    # === パスワードタブ用 デフォルト名 ===
    def get_lock_default_name(self, mode: str, src: Path) -> str:
        if mode == "set":
            return f"{src.stem}_locked.pdf"
        else:
            return f"{src.stem}_unlocked.pdf"

    def update_lock_output_placeholder(self):
        if not hasattr(self, "lock_output_entry"):
            return

        if not getattr(self, "lock_file", None):
            placeholder = "空欄:'元ファイル名'_locked.pdf / _unlocked.pdf"
        else:
            mode = self.lock_mode.get() if hasattr(self, "lock_mode") else "set"
            placeholder = self.get_lock_default_name(mode, self.lock_file)

        self.set_placeholder(self.lock_output_entry, placeholder)

    # ===== ファイル名・上書き確認（共通） =====
    def confirm_overwrite(self, path: Path) -> bool:
        """
        path が既に存在する場合、上書きして良いかを確認する。
        self.overwrite_all が True の場合は確認せずに上書き。
        戻り値: True = 実行続行, False = 中止
        """
        name = path.name

        # 不正文字チェック
        bad = [c for c in self.INVALID_FILENAME_CHARS if c in name]
        if bad:
            messagebox.showwarning(
                "警告",
                "ファイル名に使用できない文字が含まれています。\n\n"
                f"対象ファイル名: {name}\n"
                f"使用できない文字: {''.join(self.INVALID_FILENAME_CHARS)}"
            )
            return False

        # 存在しなければそのままOK
        if not path.exists():
            return True

        # 「全て上書き」がONなら聞かずにOK
        if self.overwrite_all.get():
            return True

        # 個別確認
        return messagebox.askyesno(
            "確認",
            f"{name} は既に存在します。\n\n"
            "このファイルを上書きしてもよろしいですか？"
        )

    # ===== PDF基本情報パネル =====
    def update_pdf_info(self, path: Optional[Path]):
        if not path or not path.exists():
            self.info_name.set("---")
            self.info_pages.set("---")
            self.info_size.set("---")
            self.info_path.set("---")
            return

        self.info_name.set(path.name)
        self.info_path.set(str(path))

        try:
            size_bytes = path.stat().st_size
            if size_bytes < 1024 * 1024:
                self.info_size.set(f"{size_bytes / 1024:.1f} KB")
            else:
                self.info_size.set(f"{size_bytes / (1024 * 1024):.2f} MB")
        except Exception:
            self.info_size.set("不明")

        try:
            reader = PdfReader(str(path))
            pages = len(reader.pages)
            self.info_pages.set(f"{pages} ページ")
        except Exception:
            self.info_pages.set("不明")

    # ---------------- 画面レイアウト ----------------
    def widgets(self):
        self.action_buttons = []

        # 中央エリア（メニュー or Notebook を切り替え表示）
        self.main_area = ttk.Frame(self)
        self.main_area.pack(fill="both", expand=True, padx=10, pady=(10, 5))

        # ===== メニュー画面 =====
        self.menu_frame = ttk.Frame(self.main_area)
        self.menu_frame.pack(fill="both", expand=True)

        ttk.Label(
            self.menu_frame,
            text="PDF 便利ツール メニュー",
            font=("", 14, "bold")
        ).pack(pady=(20, 10))

        btn_area = ttk.Frame(self.menu_frame)
        btn_area.pack(expand=True)

        def make_menu_button(text, row, col, cmd):
            btn = ttk.Button(btn_area, text=text, command=cmd)
            btn.grid(row=row, column=col, padx=15, pady=15,
                     ipadx=30, ipady=20, sticky="nsew")
            btn_area.grid_rowconfigure(row, weight=1)
            btn_area.grid_columnconfigure(col, weight=1)

        make_menu_button("PDF結合",           0, 0, lambda: self.show_feature("merge"))
        make_menu_button("ページ抽出／削除",   0, 1, lambda: self.show_feature("split"))
        make_menu_button("並び替え／回転",     1, 0, lambda: self.show_feature("reorder"))
        make_menu_button("PDF圧縮",           1, 1, lambda: self.show_feature("compress"))
        make_menu_button("PDF→Word/Excel変換", 2, 0, lambda: self.show_feature("convert"))
        make_menu_button("パスワード設定／解除", 2, 1, lambda: self.show_feature("password"))

        # ===== Notebook（メイン機能） ※最初は表示しない =====
        self.nb = ttk.Notebook(self.main_area)

        # ---- 各タブ作成 ----
        self.tab_merge = ttk.Frame(self.nb)
        self.tab_split = ttk.Frame(self.nb)
        self.tab_reorder = ttk.Frame(self.nb)
        self.tab_compress = ttk.Frame(self.nb)
        self.tab_convert = ttk.Frame(self.nb)      # 変換タブ
        self.tab_password = ttk.Frame(self.nb)

        self.nb.add(self.tab_merge, text="結合")
        self.nb.add(self.tab_split, text="抽出／削除")
        self.nb.add(self.tab_reorder, text="並び替え／回転")
        self.nb.add(self.tab_compress, text="圧縮")
        self.nb.add(self.tab_convert, text="変換（Word/Excel）")
        self.nb.add(self.tab_password, text="パスワード設定／解除")

        # タブ内容
        self.merge_tab()
        self.split_tab()
        self.reorder_tab()
        self.compress_tab()
        self.convert_tab()
        self.password_tab()

        # ------ 共通 出力フォルダ行 ------
        out_row = ttk.Frame(self)
        out_row.pack(fill="x", padx=10, pady=(0, 2))

        ttk.Label(out_row, text="出力フォルダ（他タブ共通）:").pack(side="left")

        out_entry = ttk.Entry(out_row, textvariable=self.output_dir_var)
        out_entry.pack(side="left", fill="x", expand=True, padx=5)

        ttk.Button(
            out_row,
            text="参照",
            command=self.browse_output_dir,
        ).pack(side="left")

        # 説明
        note_row = ttk.Frame(self)
        note_row.pack(fill="x", padx=10, pady=(0, 2))

        ttk.Label(
            note_row,
            text="※ 出力フォルダ未指定の場合、元のPDFと同じフォルダに作成します。"
        ).pack(anchor="w")

        #  処理完了後にフォルダを開くかどうか ＋ 上書き一括設定
        chk_row = ttk.Frame(self)
        chk_row.pack(fill="x", padx=10, pady=(0, 2))

        self.open_after = tk.BooleanVar(value=False)
        chk_open = ttk.Checkbutton(
            chk_row,
            text="処理完了後に出力フォルダを開く",
            variable=self.open_after,
        )
        chk_open.pack(anchor="w")
        self.action_buttons.append(chk_open)

        chk_overwrite = ttk.Checkbutton(
            chk_row,
            text="同名ファイルは確認せず全て上書きする",
            variable=self.overwrite_all,
        )
        chk_overwrite.pack(anchor="w", pady=(2, 0))
        self.action_buttons.append(chk_overwrite)

        # --- PDF基本情報パネル ---
        info_frame = ttk.LabelFrame(self, text="選択中PDFの情報")
        info_frame.pack(fill="x", padx=10, pady=(5, 5))

        row1 = ttk.Frame(info_frame)
        row1.pack(fill="x", padx=5, pady=2)
        ttk.Label(row1, text="ファイル名:").pack(side="left")
        ttk.Label(row1, textvariable=self.info_name).pack(side="left", padx=5)

        row2 = ttk.Frame(info_frame)
        row2.pack(fill="x", padx=5, pady=2)
        ttk.Label(row2, text="ページ数:").pack(side="left")
        ttk.Label(row2, textvariable=self.info_pages).pack(side="left", padx=5)

        row3 = ttk.Frame(info_frame)
        row3.pack(fill="x", padx=5, pady=2)
        ttk.Label(row3, text="ファイルサイズ:").pack(side="left")
        ttk.Label(row3, textvariable=self.info_size).pack(side="left", padx=5)

        row4 = ttk.Frame(info_frame)
        row4.pack(fill="x", padx=5, pady=2)
        ttk.Label(row4, text="場所:").pack(side="left")
        ttk.Label(row4, textvariable=self.info_path, wraplength=700).pack(side="left", padx=5)

        # --- ステータスバーとプログレスバー ---
        status_frame = ttk.Frame(self)
        status_frame.pack(fill="x", padx=5, pady=(0, 5))

        ttk.Label(
            status_frame,
            textvariable=self.status,
            anchor="w",
        ).pack(side="left", padx=5)

        # プログレスバー
        self.progress_var = tk.DoubleVar(value=0)
        self.progress = ttk.Progressbar(
            status_frame,
            variable=self.progress_var,
            maximum=100,
            mode="determinate",
            length=300,
        )
        self.progress.pack(side="right", padx=5)

    # ===== メニュー切り替え =====
    def show_menu(self):
        """Notebook を隠してメニュー画面を表示"""
        self.nb.pack_forget()
        self.menu_frame.pack(fill="both", expand=True)
        self.status.set("メニュー画面")

    def show_feature(self, feature: str):
        """メニューを隠して該当タブを表示"""
        self.menu_frame.pack_forget()
        self.nb.pack(fill="both", expand=True)

        if feature == "merge":
            self.nb.select(self.tab_merge)
        elif feature == "split":
            self.nb.select(self.tab_split)
        elif feature == "reorder":
            self.nb.select(self.tab_reorder)
        elif feature == "compress":
            self.nb.select(self.tab_compress)
        elif feature == "convert":
            self.nb.select(self.tab_convert)
        elif feature == "password":
            self.nb.select(self.tab_password)

    def browse_output_dir(self):
        initial = self.output_dir_var.get() or (
            str(self.pdf_paths[0].parent) if self.pdf_paths else ""
        )

        folder = filedialog.askdirectory(
            title="出力フォルダ（共通）を選択",
            initialdir=initial or None,
        )
        if not folder:
            return

        self.output_dir_var.set(folder)
        self.status.set(f"出力フォルダ（共通）を設定しました: {folder}")

    # ===== ボタン有効/無効ヘルパー =====
    def set_actions_state(self, enabled: bool):
        state = "normal" if enabled else "disabled"
        if hasattr(self, "action_buttons"):
            for btn in self.action_buttons:
                try:
                    btn.configure(state=state)
                except Exception:
                    pass

    # ===== プログレスバー関連ヘルパー =====
    def progress_reset(self):
        if hasattr(self, "progress_var"):
            self.progress_var.set(0)
            self.progress.update_idletasks()

    def progress_set(self, value: float):
        if hasattr(self, "progress_var"):
            self.progress_var.set(value)
            self.progress.update_idletasks()

    def progress_done(self):
        if hasattr(self, "progress_var"):
            self.progress_var.set(100)
            self.progress.update_idletasks()

    def progress_start_indeterminate(self):
        if hasattr(self, "progress"):
            self.progress.configure(mode="indeterminate", maximum=100)
            self.progress.start(10)

    def progress_stop_indeterminate(self):
        if hasattr(self, "progress"):
            self.progress.stop()
            self.progress.configure(mode="determinate", maximum=100)
            self.progress_var.set(0)
            self.progress.update_idletasks()

    # ===== D&D 共通ヘルパー =====
    def _iter_dnd_pdf_paths(self, event) -> list[Path]:
        raw = event.data
        paths = split_dnd_paths(raw)
        result = []
        for p in paths:
            path = Path(p)
            if path.suffix.lower() == ".pdf":
                result.append(path)
        return result

    # ======================= 結合タブ =======================
    def add_files(self):
        paths = filedialog.askopenfilenames(
            title="PDFファイルを選択",
            filetypes=[("PDFファイル", "*.pdf")],
        )
        if not paths:
            return

        added = 0
        for p in paths:
            path = Path(p)
            if path not in self.pdf_paths:
                self.pdf_paths.append(path)
                self.listbox.insert("end", path.name)
                added += 1

        if added > 0 and self.listbox.size() > 0:
            if not self.listbox.curselection():
                self.listbox.selection_set(0)
            self.update_merge_preview()

        self.status.set(f"{added} 件のファイルを追加しました。(合計 {len(self.pdf_paths)} 件)")
        self.update_merge_output_placeholder()

    def remove_selected(self):
        sel = list(self.listbox.curselection())
        if not sel:
            return

        for i in reversed(sel):
            self.listbox.delete(i)
            del self.pdf_paths[i]

        self.update_merge_preview()
        self.status.set(f"{len(sel)} 件のファイルを削除しました。(合計 {len(self.pdf_paths)} 件)")
        self.update_merge_output_placeholder()

    def clear_files(self):
        self.listbox.delete(0, "end")
        self.pdf_paths.clear()
        self.merge_preview_label.configure(image="", text="（PDF未選択）")
        self.merge_preview_image = None
        self.status.set("ファイルリストをクリアしました。")
        self.update_merge_output_placeholder()
        self.update_pdf_info(None)

    def move_up(self):
        sel = list(self.listbox.curselection())
        if not sel or sel[0] == 0:
            return

        for i in sel:
            if i == 0:
                continue
            self.pdf_paths[i - 1], self.pdf_paths[i] = self.pdf_paths[i], self.pdf_paths[i - 1]
            text = self.listbox.get(i)
            above = self.listbox.get(i - 1)
            self.listbox.delete(i - 1, i)
            self.listbox.insert(i - 1, text)
            self.listbox.insert(i, above)

        self.listbox.selection_clear(0, "end")
        for i in [idx - 1 for idx in sel]:
            self.listbox.selection_set(i)

        self.update_merge_preview()

    def move_down(self):
        sel = list(self.listbox.curselection())
        if not sel:
            return

        max_idx = self.listbox.size() - 1
        if sel[-1] == max_idx:
            return

        for i in reversed(sel):
            if i == max_idx:
                continue
            self.pdf_paths[i + 1], self.pdf_paths[i] = self.pdf_paths[i], self.pdf_paths[i + 1]
            text = self.listbox.get(i)
            below = self.listbox.get(i + 1)
            self.listbox.delete(i, i + 1)
            self.listbox.insert(i, below)
            self.listbox.insert(i + 1, text)

        self.listbox.selection_clear(0, "end")
        for i in [idx + 1 for idx in sel]:
            self.listbox.selection_set(i)

        self.update_merge_preview()

    def on_drop_merge(self, event):
        pdf_paths = self._iter_dnd_pdf_paths(event)
        if not pdf_paths:
            return

        added = 0
        for path in pdf_paths:
            if path not in self.pdf_paths:
                self.pdf_paths.append(path)
                self.listbox.insert("end", path.name)
                added += 1

        if added > 0 and self.listbox.size() > 0:
            if not self.listbox.curselection():
                self.listbox.selection_set(0)
            self.update_merge_preview()

        self.status.set(f"D&Dで {added} 件追加しました（合計 {len(self.pdf_paths)} 件）")
        self.update_merge_output_placeholder()

    def merge_tab(self):
        frame = self.tab_merge

        # ← メニューに戻る
        top_bar = ttk.Frame(frame)
        top_bar.pack(fill="x", padx=10, pady=(10, 0))
        ttk.Button(top_bar, text="← メニューに戻る", command=self.show_menu).pack(side="right")

        ttk.Label(frame, text="選択中のPDFを、この順番で結合します。").pack(anchor="w", padx=10, pady=10)
        ttk.Label(
            frame,
            text="PDFを追加ボタンから追加してください。ドラッグ＆ドロップでも追加可能です。"
        ).pack(anchor="w", padx=10, pady=10)

        mid = ttk.Frame(frame)
        mid.pack(fill="both", expand=True, padx=10, pady=(5, 5))

        # 左：PDF一覧
        list_area = ttk.Frame(mid)
        list_area.pack(side="left", fill="both", expand=True)

        list_frame = ttk.Frame(list_area, height=140)
        list_frame.pack(side="left", fill="both", expand=True)
        list_frame.pack_propagate(False)

        self.listbox = tk.Listbox(list_frame, height=6, width=50)
        self.listbox.pack(side="left", fill="both", expand=True)

        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.listbox.yview)
        scrollbar.pack(side="right", fill="y")
        self.listbox.config(yscrollcommand=scrollbar.set)

        self.listbox.bind("<<ListboxSelect>>", self.update_merge_preview)

        if DND_AVAILABLE:
            self.listbox.drop_target_register(DND_FILES)
            self.listbox.dnd_bind("<<Drop>>", self.on_drop_merge)

        btn_frame = ttk.Frame(list_area)
        btn_frame.pack(side="left", fill="y", padx=(10, 0))

        btn_add = ttk.Button(btn_frame, text="追加", command=self.add_files)
        btn_add.pack(fill="x", pady=2)
        self.action_buttons.append(btn_add)

        btn_del = ttk.Button(btn_frame, text="削除", command=self.remove_selected)
        btn_del.pack(fill="x", pady=2)
        self.action_buttons.append(btn_del)

        btn_clear = ttk.Button(btn_frame, text="クリア", command=self.clear_files)
        btn_clear.pack(fill="x", pady=2)
        self.action_buttons.append(btn_clear)

        ttk.Separator(btn_frame, orient="horizontal").pack(fill="x", pady=4)

        btn_up = ttk.Button(btn_frame, text="上へ", command=self.move_up)
        btn_up.pack(fill="x", pady=2)
        self.action_buttons.append(btn_up)

        btn_down = ttk.Button(btn_frame, text="下へ", command=self.move_down)
        btn_down.pack(fill="x", pady=2)
        self.action_buttons.append(btn_down)

        # 右：プレビュー
        preview_group = ttk.LabelFrame(mid, text="プレビュー（選択中PDFの1ページ目）")
        preview_group.pack(side="left", fill="both", expand=True, padx=(10, 0))

        self.merge_preview_label = ttk.Label(preview_group, anchor="center")
        self.merge_preview_label.pack(fill="both", expand=True, padx=5, pady=5)
        self.merge_preview_label.configure(text="（PDF未選択）")

        # 出力ファイル名
        out_frame = ttk.Frame(frame)
        out_frame.pack(fill="x", padx=10, pady=(0, 10))

        ttk.Label(out_frame, text="出力ファイル名:").pack(side="left")

        self.merge_output_var = tk.StringVar(value="")
        self.merge_output_entry = ttk.Entry(out_frame, textvariable=self.merge_output_var, width=50)
        self.merge_output_entry.pack(side="left", padx=5)

        self.init_placeholder(self.merge_output_entry, self.get_merge_default_name())

        btn_run_merge = ttk.Button(frame, text="結合を実行", command=self.run_merge)
        btn_run_merge.pack(pady=10)
        self.action_buttons.append(btn_run_merge)

    def update_merge_preview(self, event=None):
        if not hasattr(self, "merge_preview_label"):
            return

        if not self.pdf_paths or self.listbox.size() == 0:
            self.merge_preview_label.configure(image="", text="（PDF未選択）")
            self.merge_preview_image = None
            self.update_pdf_info(None)
            return

        sel = self.listbox.curselection()
        if not sel:
            idx = 0
        else:
            idx = sel[0]

        if not (0 <= idx < len(self.pdf_paths)):
            return

        path = self.pdf_paths[idx]
        self.update_pdf_info(path)

        target_h = 380

        try:
            doc = pdfium.PdfDocument(str(path))
            page = doc[0]

            w_pt, h_pt = page.get_size()
            if h_pt == 0:
                scale = 1.0
            else:
                scale = target_h / h_pt

            if scale < 0.2:
                scale = 0.2
            if scale > 3.0:
                scale = 3.0

            pil_image = page.render(scale=scale).to_pil()
            img = ImageTk.PhotoImage(pil_image)

            self.merge_preview_image = img
            self.merge_preview_label.configure(image=img, text="")

            doc.close()

        except Exception as e:
            self.merge_preview_label.configure(
                image="",
                text=f"プレビューに失敗しました。\n{path.name}\n{e}"
            )
            self.merge_preview_image = None

    @staticmethod
    def merge_pdfs(inputs, output, progress_cb=None):
        writer = PdfWriter()
        total = len(inputs)

        for i, p in enumerate(inputs, start=1):
            path = Path(p)
            reader = PdfReader(str(p))
            if reader.is_encrypted:
                raise ValueError(
                    f"パスワードが設定されているPDFは結合できません。\n"
                    f"対象ファイル: {path.name}"
                )

            for page in reader.pages:
                writer.add_page(page)

            if progress_cb is not None and total > 0:
                percent = (i / total) * 100
                progress_cb(percent)

        output.parent.mkdir(parents=True, exist_ok=True)
        with open(output, "wb") as f:
            writer.write(f)

    def run_merge(self):
        if len(self.pdf_paths) < 2:
            messagebox.showwarning("警告", "結合するには2つ以上のPDFを選択してください。")
            return

        out_name = self.get_entry_text(self.merge_output_entry)
        if not out_name:
            out_name = self.get_merge_default_name()

        if not out_name.lower().endswith(".pdf"):
            out_name += ".pdf"

        dir_str = self.output_dir_var.get().strip()
        if dir_str:
            out_dir = Path(dir_str)
        else:
            if not self.pdf_paths:
                messagebox.showwarning("警告", "PDFファイルが選択されていません。")
                return
            out_dir = self.pdf_paths[0].parent

        if not out_dir.exists():
            try:
                out_dir.mkdir(parents=True, exist_ok=True)
            except Exception as e:
                messagebox.showerror("エラー", f"出力フォルダの作成に失敗しました:\n{e}")
                return

        out_path = out_dir / out_name

        if not self.confirm_overwrite(out_path):
            self.status.set("結合をキャンセルしました（既存ファイルあり）")
            return

        self.progress_reset()
        self.status.set("結合処理を開始しました")
        self.set_actions_state(False)
        try:
            self.merge_pdfs(self.pdf_paths, out_path, progress_cb=self.progress_set)
        except ValueError as e:
            messagebox.showwarning("警告", str(e))
            self.status.set("結合を中止しました（パスワード付きPDFが含まれています）")
        except Exception as e:
            messagebox.showerror("エラー", f"結合中にエラーが発生しました:\n{e}")
            self.status.set("結合に失敗しました")
        else:
            messagebox.showinfo("完了", f"結合しました:\n{out_path}")
            self.progress_done()
            self.status.set(f"結合を完了しました: {out_path}")
            if self.open_after.get():
                open_folder(out_path)
        finally:
            self.set_actions_state(True)

    # ======================= 抽出／削除タブ =======================
    def split_tab(self):
        frame = self.tab_split

        # ← メニューに戻る
        top_bar = ttk.Frame(frame)
        top_bar.pack(fill="x", padx=10, pady=(10, 0))
        ttk.Button(top_bar, text="← メニューに戻る", command=self.show_menu).pack(side="right")

        ttk.Label(
            frame,
            text=(
                "選択したPDFからページを抽出/削除します。\nPDFを選択してください。ドラッグ＆ドロップでも追加可能です。\n\n"
                "サムネイルをクリックして,抽出or削除したいページを選択するか、ページ指定欄に直接入力してください。"
            )
        ).pack(anchor="w", padx=10, pady=10)

        src_frame = ttk.Frame(frame)
        src_frame.pack(fill="x", padx=10, pady=(0, 10))

        self.split_src_path: Optional[Path] = None
        self.split_src_label = tk.StringVar(value="(未選択)")

        btn_choose = ttk.Button(
            src_frame,
            text="PDFを選択",
            command=self.choose_split_pdf,
        )
        btn_choose.pack(side="left")
        self.action_buttons.append(btn_choose)

        btn_clear = ttk.Button(
            src_frame,
            text="クリア",
            command=self.clear_split_pdf,
        )
        btn_clear.pack(side="left", padx=5)
        self.action_buttons.append(btn_clear)

        ttk.Label(
            src_frame,
            textvariable=self.split_src_label,
        ).pack(side="left", padx=10)

        thumb_frame = ttk.LabelFrame(frame, text="ページサムネイル（クリックで選択／Ctrl+クリックで複数選択）")
        thumb_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        self.split_thumb_view = PageSelectView(thumb_frame, thumb_height=90)
        self.split_thumb_view.pack(fill="both", expand=True)

        if DND_AVAILABLE:
            thumb_frame.drop_target_register(DND_FILES)
            thumb_frame.dnd_bind("<<Drop>>", self.on_drop_split)

        range_frame = ttk.Frame(frame)
        range_frame.pack(fill="x", padx=10, pady=(0, 5))

        ttk.Label(range_frame, text="ページ指定(例: 1,3,5-7):").pack(side="left")
        self.page_range_var = tk.StringVar(value="")
        ttk.Entry(range_frame, textvariable=self.page_range_var).pack(
            side="left", fill="x", expand=True, padx=5
        )

        ttk.Label(
            frame,
            text="※ テキスト指定が空欄の場合、サムネイルで選択したページが対象になります。"
        ).pack(anchor="w", padx=10)

        mode_frame = ttk.Frame(frame)
        mode_frame.pack(fill="x", padx=10, pady=(5, 5))

        ttk.Label(mode_frame, text="処理モード").pack(side="left")

        self.page_edit_mode = tk.StringVar(value="keep")
        ttk.Radiobutton(
            mode_frame, text="抽出（指定ページだけ残す）",
            variable=self.page_edit_mode, value="keep",
            command=self.update_split_output_placeholder,
        ).pack(side="left", padx=5)
        ttk.Radiobutton(
            mode_frame, text="削除（指定ページを消す）",
            variable=self.page_edit_mode, value="delete",
            command=self.update_split_output_placeholder,
        ).pack(side="left", padx=5)

        out_frame = ttk.Frame(frame)
        out_frame.pack(fill="x", padx=10, pady=(0, 10))

        ttk.Label(out_frame, text="出力ファイル名:").pack(side="left")

        self.split_output_var = tk.StringVar(value="")
        self.split_output_entry = ttk.Entry(out_frame, textvariable=self.split_output_var, width=50)
        self.split_output_entry.pack(side="left", padx=5)

        self.init_placeholder(
            self.split_output_entry,
            "空欄:'元ファイル名'_抽出済み.pdf"
        )

        btn_run_split = ttk.Button(frame, text="ページ抽出/削除を実行", command=self.run_split)
        btn_run_split.pack(pady=10)
        self.action_buttons.append(btn_run_split)

    def clear_split_pdf(self):
        self.split_src_path = None
        self.split_src_label.set("(未選択)")
        if hasattr(self, "split_thumb_view"):
            try:
                self.split_thumb_view.clear()
            except Exception:
                pass
        self.update_pdf_info(None)
        self.update_split_output_placeholder()
        self.status.set("抽出／削除対象をクリアしました。")

    def on_drop_split(self, event):
        pdf_paths = self._iter_dnd_pdf_paths(event)
        if not pdf_paths:
            return

        path = pdf_paths[0]
        self.split_src_path = path
        self.split_src_label.set(path.name)
        self.status.set(f"抽出／削除対象（D&D）: {path}")
        self.update_pdf_info(path)

        if hasattr(self, "split_thumb_view"):
            try:
                self.split_thumb_view.load_pdf(str(path))
            except Exception as e:
                messagebox.showerror("エラー", f"サムネイル作成に失敗しました:\n{e}")

        self.update_split_output_placeholder()

    def choose_split_pdf(self):
        path = filedialog.askopenfilename(
            title="抽出／削除するPDFを選択",
            filetypes=[("PDFファイル", "*.pdf")],
        )
        if not path:
            return

        self.split_src_path = Path(path)
        self.split_src_label.set(self.split_src_path.name)
        self.status.set(f"抽出／削除対象: {self.split_src_path}")
        self.update_pdf_info(self.split_src_path)

        if hasattr(self, "split_thumb_view"):
            try:
                self.split_thumb_view.load_pdf(str(self.split_src_path))
            except Exception as e:
                messagebox.showerror("エラー", f"サムネイル作成に失敗しました:\n{e}")

        self.update_split_output_placeholder()

    def run_split(self):
        if not getattr(self, "split_src_path", None):
            messagebox.showwarning("警告", "抽出／削除するPDFを選択してください。")
            return

        src_path = self.split_src_path

        dir_str = self.output_dir_var.get().strip()
        if dir_str:
            out_dir = Path(dir_str)
        else:
            out_dir = src_path.parent

        try:
            out_dir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            messagebox.showerror("エラー", f"出力フォルダの作成に失敗しました:\n{e}")
            return

        text = self.page_range_var.get().strip()

        try:
            reader = PdfReader(str(src_path))
        except Exception as e:
            messagebox.showerror("エラー", f"PDFの読み込みに失敗しました:\n{e}")
            return

        if reader.is_encrypted:
            messagebox.showwarning(
                "警告",
                "このPDFにはパスワードが設定されているため、\n"
                "抽出／削除タブでは処理できません。\n\n"
            )
            return

        total_pages = len(reader.pages)
        if total_pages == 0:
            messagebox.showwarning("警告", "PDFにページが含まれていません。")
            return

        target_indices: list[int] = []

        if text:
            try:
                target_indices = self._parse_page_ranges(text, total_pages)
                if not target_indices:
                    messagebox.showwarning("警告", "有効なページ番号がありません。")
                    return
            except ValueError as e:
                messagebox.showwarning("警告", f"ページ指定の形式が不正です:\n{e}")
                return
        else:
            if hasattr(self, "split_thumb_view"):
                target_indices = self.split_thumb_view.get_selected_indices()
            if not target_indices:
                messagebox.showwarning(
                    "警告",
                    "ページ指定がありません。\nテキストでページを指定するか、サムネイルを選択してください。"
                )
                return

        mode = self.page_edit_mode.get()

        if mode == "keep":
            kept = target_indices
        else:
            to_delete = set(target_indices)
            kept = [i for i in range(total_pages) if i not in to_delete]

        if not kept:
            messagebox.showwarning("警告", "結果として残るページがありません。")
            return

        writer = PdfWriter()
        for idx in kept:
            writer.add_page(reader.pages[idx])

        raw_name = self.get_entry_text(self.split_output_entry)
        if not raw_name:
            raw_name = self.get_split_default_name(mode, src_path)

        if not raw_name.lower().endswith(".pdf"):
            raw_name += ".pdf"

        out_path = out_dir / raw_name

        if not self.confirm_overwrite(out_path):
            self.status.set(f"ページ{'抽出' if mode == 'keep' else '削除'}をキャンセルしました（既存ファイルあり）")
            return

        self.set_actions_state(False)

        try:
            with open(out_path, "wb") as f:
                writer.write(f)
        except Exception as e:
            messagebox.showerror("エラー", f"書き出し中にエラーが発生しました:\n{e}")
            self.status.set("抽出／削除に失敗しました")
            self.set_actions_state(True)
            return

        self.set_actions_state(True)

        mode_jp = "抽出" if mode == "keep" else "削除"
        messagebox.showinfo(
            "完了",
            f"{mode_jp}を完了しました。\n出力ファイル:{out_path}\n"
            f"元ページ数:{total_pages} / 残りページ数:{len(kept)}"
        )
        self.status.set(f"ページ{mode_jp}を完了しました:{out_path}")
        if self.open_after.get():
            open_folder(out_path)

    def _parse_page_ranges(self, text: str, total_pages: int) -> list[int]:
        result: set[int] = set()

        parts = text.split(",")
        for part in parts:
            part = part.strip()
            if not part:
                continue

            if "-" in part:
                start_str, end_str = part.split("-", 1)
                start = int(start_str)
                end = int(end_str)
                if start < 1 or end < 1 or start > end:
                    raise ValueError(f"範囲指定が不正です: {part}")
                for p in range(start, end + 1):
                    if p > total_pages:
                        raise ValueError(f"ページ番号 {p} は {total_pages} ページを超えています。")
                    result.add(p - 1)
            else:
                p = int(part)
                if p < 1 or p > total_pages:
                    raise ValueError(f"ページ番号 {p} は 1〜{total_pages} の範囲外です。")
                result.add(p - 1)

        return sorted(result)

    # ======================= 並び替え／回転タブ =======================
    def reorder_tab(self):
        frame = self.tab_reorder

        # ← メニューに戻る
        top_bar = ttk.Frame(frame)
        top_bar.pack(fill="x", padx=10, pady=(10, 0))
        ttk.Button(top_bar, text="← メニューに戻る", command=self.show_menu).pack(side="right")

        ttk.Label(
            frame,
            text=(
                "選択したPDFのページ順を入れ替えます。\nPDFを選択してください。ドラッグ＆ドロップにて入れ替えることができます。\n"
                "下の回転ボタンで、選択中のページを回転させることができます。"
            ),
        ).pack(anchor="w", padx=10, pady=10)

        order_frame = ttk.Frame(frame)
        order_frame.pack(fill="x", padx=10, pady=(0, 10))

        btn_choose = ttk.Button(order_frame, text="PDFを選択", command=self.choose_reorder_pdf)
        btn_choose.pack(side="left")
        self.action_buttons.append(btn_choose)

        btn_clear = ttk.Button(order_frame, text="クリア", command=self.clear_reorder_pdf)
        btn_clear.pack(side="left", padx=5)
        self.action_buttons.append(btn_clear)

        self.reorder_var = tk.StringVar(value="(未選択)")
        ttk.Label(order_frame, textvariable=self.reorder_var).pack(side="left", padx=10)

        thumb_group = ttk.LabelFrame(
            frame,
            text="ページサムネイル（ドラッグ＆ドロップで並び替え／クリックでプレビュー／Ctrl+クリックで複数選択）"
        )
        thumb_group.pack(fill="both", expand=True, padx=10, pady=5)

        self.reorder_thumb_view = PageThumbnailView(thumb_group, thumb_height=90)
        self.reorder_thumb_view.pack(fill="both", expand=True, padx=5, pady=5)

        if DND_AVAILABLE:
            thumb_group.drop_target_register(DND_FILES)
            thumb_group.dnd_bind("<<Drop>>", self.on_drop_reorder)

        rotate_frame = ttk.Frame(frame)
        rotate_frame.pack(fill="x", padx=10, pady=(5, 10))

        ttk.Label(rotate_frame, text="選択ページを回転:").pack(side="left")
        ttk.Button(
            rotate_frame,
            text="⟲ 90°左",
            command=lambda: self.reorder_thumb_view.rotate_selected(-90),
        ).pack(side="left", padx=5)

        ttk.Button(
            rotate_frame,
            text="⟳ 90°右",
            command=lambda: self.reorder_thumb_view.rotate_selected(90),
        ).pack(side="left", padx=5)

        out_frame = ttk.Frame(frame)
        out_frame.pack(fill="x", padx=10, pady=(5, 5))

        ttk.Label(out_frame, text="出力ファイル名:").pack(side="left")

        self.reorder_output_var = tk.StringVar(value="")
        self.reorder_output_entry = ttk.Entry(out_frame, textvariable=self.reorder_output_var, width=50)
        self.reorder_output_entry.pack(side="left", padx=5)

        self.init_placeholder(
            self.reorder_output_entry,
            "空欄:'元ファイル名'_並び替え済み.pdf"
        )

        btn_run_reorder = ttk.Button(
            frame, text="並び替え/回転を実行", command=self.run_reorder
        )
        btn_run_reorder.pack(pady=(0, 10))
        self.action_buttons.append(btn_run_reorder)

    def clear_reorder_pdf(self):
        self.reorder_pdf_path = None
        self.reorder_var.set("(未選択)")
        if hasattr(self, "reorder_thumb_view"):
            try:
                self.reorder_thumb_view.clear()
            except Exception:
                pass
        self.update_pdf_info(None)
        self.update_reorder_output_placeholder()
        self.status.set("並び替え対象をクリアしました。")

    def on_drop_reorder(self, event):
        pdf_paths = self._iter_dnd_pdf_paths(event)
        if not pdf_paths:
            return

        path = pdf_paths[0]
        self.reorder_pdf_path = str(path)
        self.reorder_var.set(path.name)
        self.status.set(f"並び替え対象（D&D）:{path}")
        self.update_pdf_info(path)

        try:
            self.reorder_thumb_view.load_pdf(str(path))
        except Exception as e:
            messagebox.showerror(
                "エラー", f"サムネイル作成中にエラーが発生しました:\n{e}"
            )
            self.status.set("サムネイル作成に失敗しました")

        self.update_reorder_output_placeholder()

    def run_reorder(self):
        if not getattr(self, "reorder_pdf_path", None):
            messagebox.showerror("警告", "並び替え対象のPDFが選択されていません。")
            return

        in_path = Path(self.reorder_pdf_path)

        dir_str = self.output_dir_var.get().strip()
        if dir_str:
            out_dir = Path(dir_str)
        else:
            out_dir = in_path.parent

        raw_name = self.get_entry_text(self.reorder_output_entry)
        if not raw_name:
            raw_name = self.get_reorder_default_name(in_path)

        if not raw_name.lower().endswith(".pdf"):
            raw_name += ".pdf"

        out_path = out_dir / raw_name

        if not self.confirm_overwrite(out_path):
            self.status.set("並び替え/回転をキャンセルしました（既存ファイルあり）")
            return

        order = self.reorder_thumb_view.get_page_order()
        if not order:
            messagebox.showwarning("警告", "ページ情報が取得できません。")
            return

        rotations = self.reorder_thumb_view.get_page_rotations()

        reader = PdfReader(str(in_path))

        if reader.is_encrypted:
            messagebox.showwarning(
                "警告",
                "このPDFにはパスワードが設定されているため、\n"
                "並び替え／回転タブでは処理できません。\n\n"
            )
            return

        writer = PdfWriter()

        total = len(order)
        self.progress_reset()
        self.status.set("ページ並べ替え/回転中...")
        self.set_actions_state(False)

        try:
            for i, src_idx in enumerate(order, start=1):
                page = reader.pages[src_idx]

                angle = rotations.get(src_idx, 0)
                if angle:
                    try:
                        page.rotate(angle)
                    except Exception:
                        try:
                            page = page.rotate_clockwise(angle)
                        except Exception as e:
                            messagebox.showerror("エラー", f"ページ回転中にエラーが発生しました。\n{e}")
                            self.set_actions_state(True)
                            return

                writer.add_page(page)

                percent = i / total * 100
                self.progress_set(percent)

            out_path.parent.mkdir(parents=True, exist_ok=True)
            with open(out_path, "wb") as f:
                writer.write(f)

        except Exception as e:
            messagebox.showerror("エラー", f"並び替え/回転中にエラーが発生しました:\n{e}")
            self.status.set("並び替え/回転に失敗しました")
            self.set_actions_state(True)
            return

        self.set_actions_state(True)
        self.progress_done()
        messagebox.showinfo("完了", f"並び替え・回転を完了しました:\n{out_path}")
        self.status.set(f"並び替え・回転を完了しました: {out_path}")
        if self.open_after.get():
            open_folder(out_path)

    def choose_reorder_pdf(self):
        path = filedialog.askopenfilename(
            title="並び替えをするPDFを選択",
            filetypes=[("PDFファイル", "*.pdf")]
        )
        if not path:
            return

        self.reorder_pdf_path = path
        p = Path(path)
        self.reorder_var.set(p.name)
        self.update_pdf_info(p)
        self.status.set(f"並び替え対象:{path}")

        try:
            self.reorder_thumb_view.load_pdf(path)
        except Exception as e:
            messagebox.showerror(
                "エラー", f"サムネイル作成中にエラーが発生しました:\n{e}"
            )
            self.status.set("サムネイル作成に失敗しました")

        self.update_reorder_output_placeholder()

    # ======================= パスワード関連（ヘルパー） =======================
    @staticmethod
    def set_pdf_password(src: Path, out_path: Path, owner_password: str, allow_print: bool) -> None:
        """
        閲覧は自由（パスワード不要）だが、編集・注釈・フォーム入力などを禁止する PDF を出力する。
        """
        reader = PdfReader(str(src))
        writer = PdfWriter()

        for page in reader.pages:
            writer.add_page(page)

        # いったん「全許可」からスタート
        perms = UserAccessPermissions(-1)

        # 編集・組み立て・注釈などを禁止
        deny_flags = [
            "MODIFY",
            "ASSEMBLE_DOC",
            "ADD_OR_MODIFY",
            "FILL_FORM_FIELDS",
            "EXTRACT_TEXT_AND_GRAPHICS",
        ]

        for name in deny_flags:
            flag = getattr(UserAccessPermissions, name, None)
            if flag is not None:
                perms &= ~flag

        # 印刷禁止指定
        if not allow_print:
            flag_print = getattr(UserAccessPermissions, "PRINT", None)
            if flag_print is not None:
                perms &= ~flag_print

        writer.encrypt(
            user_password="",
            owner_password=owner_password,
            permissions_flag=perms,
        )

        out_path.parent.mkdir(parents=True, exist_ok=True)
        with open(out_path, "wb") as f:
            writer.write(f)

    @staticmethod
    def remove_pdf_password(src: Path, out_path: Path, password: str) -> None:
        reader = PdfReader(str(src))

        if not reader.is_encrypted:
            raise ValueError("このPDFにはパスワードが設定されていません。")

        result = reader.decrypt(password)
        if result == 0:
            raise ValueError("パスワードが違います。")

        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)

        out_path.parent.mkdir(parents=True, exist_ok=True)
        with out_path.open("wb") as f:
            writer.write(f)

    # ======================= 圧縮タブ =======================
    @staticmethod
    def compress_one_pdf(input_path: Path, output_path: Path, gs_path: str, pdf_settings: str) -> None:
        cmd = [
            gs_path,
            "-sDEVICE=pdfwrite",
            "-dCompatibilityLevel=1.4",
            f"-dPDFSETTINGS={pdf_settings}",
            "-dNOPAUSE",
            "-dQUIET",
            "-dBATCH",
            f"-sOutputFile={str(output_path)}",
            str(input_path),
        ]

        if os.name == "nt":
            CREATE_NO_WINDOW = 0x08000000
            subprocess.run(cmd, check=True, creationflags=CREATE_NO_WINDOW)
        else:
            subprocess.run(cmd, check=True)

    def compress_tab(self):
        frame = self.tab_compress

        # ← メニューに戻る
        top_bar = ttk.Frame(frame)
        top_bar.pack(fill="x", padx=10, pady=(10, 0))
        ttk.Button(top_bar, text="← メニューに戻る", command=self.show_menu).pack(side="right")

        targets_frame = ttk.Frame(frame)
        targets_frame.pack(fill="x", padx=10, pady=(10, 10))

        ttk.Label(targets_frame, text="圧縮するPDF（複数選択可）:").pack(side="left")

        self.compress_files: list[Path] = []
        self.compress_label = tk.StringVar(value="（未選択）")

        btn_choose = ttk.Button(
            targets_frame,
            text="PDFを選択",
            command=self.choose_compress_pdfs,
        )
        btn_choose.pack(side="left")
        self.action_buttons.append(btn_choose)

        btn_clear = ttk.Button(
            targets_frame,
            text="クリア",
            command=self.clear_compress_pdfs,
        )
        btn_clear.pack(side="left", padx=5)
        self.action_buttons.append(btn_clear)

        ttk.Label(
            targets_frame,
            textvariable=self.compress_label,
        ).pack(side="left", padx=10)

        if DND_AVAILABLE:
            frame.drop_target_register(DND_FILES)
            frame.dnd_bind("<<Drop>>", self.on_drop_compress)

        row = ttk.Frame(frame)
        row.pack(fill="x", padx=10, pady=(5, 0))

        ttk.Label(
            row,
            text="圧縮レベル   低圧縮（高画質） ←→ 高圧縮(低画質）"
        ).pack(anchor="w")

        ttk.Label(
            frame,
            text="※ 数字が小さいほど画質優先／大きいほどファイルサイズ優先になります。"
        ).pack(anchor="w", padx=10)

        scale_label_frame = ttk.Frame(frame)
        scale_label_frame.pack(fill="x", padx=10)

        ttk.Label(scale_label_frame, text="1").pack(side="left")
        ttk.Label(scale_label_frame, text="3").pack(side="left", expand=True)
        ttk.Label(scale_label_frame, text="5").pack(side="right")

        self.compress_level = tk.IntVar(value=3)

        self.compress_scale = ttk.Scale(
            frame,
            from_=1,
            to=5,
            orient="horizontal",
            variable=self.compress_level,
        )
        self.compress_scale.pack(fill="x", padx=10, pady=5)
        self.compress_scale.bind("<ButtonRelease-1>", self.on_compress_scale_release)

        suffix_frame = ttk.Frame(frame)
        suffix_frame.pack(fill="x", padx=10, pady=(5, 5))

        ttk.Label(suffix_frame, text="出力ファイル名:").pack(side="left")

        self.compress_suffix_var = tk.StringVar(value="")
        self.compress_suffix_entry = ttk.Entry(
            suffix_frame,
            textvariable=self.compress_suffix_var,
            width=50,
        )
        self.compress_suffix_entry.pack(side="left", padx=5)

        self.init_placeholder(
            self.compress_suffix_entry,
            self.get_compress_default_suffix()
        )
        self.update_compress_suffix_placeholder()

        size_frame = ttk.Frame(frame)
        size_frame.pack(fill="x", padx=10, pady=(5, 10))

        ttk.Label(size_frame, text="目標サイズ（参考値）:").pack(side="left")

        self.target_size_mb = tk.StringVar(value="")
        self.target_size_entry = ttk.Entry(
            size_frame,
            textvariable=self.target_size_mb,
            width=10,
        )
        self.target_size_entry.pack(side="left", padx=5)

        ttk.Label(size_frame, text="MB").pack(side="left")

        ttk.Label(
            frame,
            text=(
                "※ 目標サイズ(MB)は「このサイズ以下を目指す上限値」です。\n"
                "   PDFの内容により、指定した値に一致するとは限らず、"
                "   それより小さくなったり、十分に小さくできない場合もあります。\n"
                "   （場合によってはサイズが大きくなる場合もあります。）"
            ),
            foreground="#444444",
            wraplength=600
        ).pack(anchor="w", padx=10, pady=(0, 10))

        btn_run_compress = ttk.Button(frame, text="圧縮を実行", command=self.run_compress)
        btn_run_compress.pack(pady=10)
        self.action_buttons.append(btn_run_compress)

    def clear_compress_pdfs(self):
        self.compress_files = []
        self.compress_label.set("（未選択）")
        self.update_compress_suffix_placeholder()
        self.update_pdf_info(None)
        self.status.set("圧縮対象をクリアしました。")

    def on_compress_scale_release(self, event=None):
        try:
            val = float(self.compress_level.get())
        except Exception:
            val = 3.0

        val = round(val)
        if val < 1:
            val = 1
        elif val > 5:
            val = 5

        self.compress_level.set(val)

    def on_drop_compress(self, event):
        pdf_paths = self._iter_dnd_pdf_paths(event)
        if not pdf_paths:
            return

        added = 0
        for path in pdf_paths:
            if path not in self.compress_files:
                self.compress_files.append(path)
                added += 1

        if not self.compress_files:
            self.compress_label.set("（未選択）")
        elif len(self.compress_files) == 1:
            self.compress_label.set(self.compress_files[0].name)
        else:
            self.compress_label.set(
                f"{self.compress_files[0].name} ほか {len(self.compress_files) - 1}件"
            )

        self.status.set(f"D&Dで圧縮対象: {len(self.compress_files)} 件のPDFを選択しました。")
        self.update_compress_suffix_placeholder()

        if self.compress_files:
            self.update_pdf_info(self.compress_files[0])

    def choose_compress_pdfs(self):
        paths = filedialog.askopenfilenames(
            title="圧縮するPDFを選択",
            filetypes=[("PDFファイル", "*.pdf")],
        )
        if not paths:
            return

        self.compress_files = [Path(p) for p in paths]

        if len(self.compress_files) == 1:
            label = self.compress_files[0].name
        else:
            label = f"{self.compress_files[0].name} ほか {len(self.compress_files) - 1}件"

        self.compress_label.set(label)
        self.status.set(f"圧縮対象: {len(self.compress_files)} 件のPDFを選択しました。")

        self.update_compress_suffix_placeholder()

        if self.compress_files:
            self.update_pdf_info(self.compress_files[0])

    def run_compress(self):
        if not getattr(self, "compress_files", None):
            self.compress_files = []

        if not self.compress_files:
            messagebox.showwarning("警告", "圧縮するPDFを選んでください。")
            return

        if not hasattr(self, "gs_path") or self.gs_path is None:
            self.gs_path = find_gs()

        if not self.gs_path:
            messagebox.showerror("エラー", "Ghostscript が見つかりません。圧縮を実行できません。")
            return

        target_mb: Optional[float] = None

        raw = self.target_size_mb.get()
        raw_stripped = raw.strip()

        if raw_stripped:
            try:
                v = float(raw_stripped)
                if v > 0:
                    target_mb = v
            except ValueError:
                messagebox.showwarning("警告", "目標サイズ(MB) は数値で入力してください。")
                return

        presets = ["/prepress", "/printer", "/default", "/ebook", "/screen"]

        def level_to_start_index(lv: int) -> int:
            if lv <= 1:
                return 0
            elif lv == 2:
                return 1
            elif lv == 3:
                return 2
            elif lv == 4:
                return 3
            else:
                return 4

        if target_mb is None:
            level = int(self.compress_level.get())
            start_idx = level_to_start_index(level)
        else:
            start_idx = 0

        tasks: list[tuple[Path, Path]] = []
        report_lines: list[str] = []

        for src in self.compress_files:
            src = Path(src)

            pattern = self.get_entry_text(self.compress_suffix_entry).strip()

            if not pattern:
                out_name = self.get_compress_default_name(src)
            else:
                if "{name}" in pattern:
                    out_name = pattern.replace("{name}", src.stem)
                else:
                    out_name = pattern

                if not out_name.lower().endswith(".pdf"):
                    out_name += ".pdf"

            dir_str = self.output_dir_var.get().strip()
            if dir_str:
                out_path = Path(dir_str) / out_name
            else:
                out_path = src.with_name(out_name)

            if not self.confirm_overwrite(out_path):
                report_lines.append(f"{src.name}: スキップ（既存ファイル有りのため）")
            else:
                tasks.append((src, out_path))

        if not tasks:
            self.status.set("圧縮をキャンセルしました（すべてスキップ）")
            return

        self.progress_reset()
        total_files = len(tasks)
        self.set_actions_state(False)
        self.status.set("圧縮処理を開始しました...")

        gs_path = self.gs_path

        def worker():
            processed = 0
            last_out_path: Optional[Path] = None

            try:
                for idx, (src, out_path) in enumerate(tasks, start=1):
                    try:
                        orig_bytes = src.stat().st_size
                    except FileNotFoundError:
                        report_lines.append(f"{src.name}: ファイルが見つかりません（スキップ）")
                        continue

                    orig_mb = orig_bytes / (1024 * 1024)

                    if target_mb is None:
                        setting = presets[start_idx]
                        self.compress_one_pdf(src, out_path, gs_path, setting)
                        new_bytes = out_path.stat().st_size
                        new_mb = new_bytes / (1024 * 1024)
                    else:
                        cur_idx = start_idx
                        last_idx = len(presets) - 1

                        while True:
                            setting = presets[cur_idx]
                            self.compress_one_pdf(src, out_path, gs_path, setting)
                            new_bytes = out_path.stat().st_size
                            new_mb = new_bytes / (1024 * 1024)

                            if new_mb <= target_mb:
                                break

                            if cur_idx < last_idx:
                                cur_idx += 1
                            else:
                                break

                    processed += 1
                    last_out_path = out_path

                    if orig_bytes > 0:
                        reduce_ratio = (1 - new_bytes / orig_bytes) * 100
                    else:
                        reduce_ratio = 0.0

                    report_lines.append(
                        f"{src.name}: {orig_mb:.2f}MB → {new_mb:.2f}MB ({reduce_ratio:.1f}%削減)"
                    )

                    percent = idx / total_files * 100

                    def _update_progress(p=percent, name=src.name, i=idx):
                        self.progress_set(p)
                        self.status.set(f"圧縮中...（{i}/{total_files}）{name}")

                    self.after(0, _update_progress)

            except subprocess.CalledProcessError as e:
                def _on_error():
                    messagebox.showerror("エラー", f"圧縮中にエラーが発生しました:\n{e}")
                    self.status.set("圧縮に失敗しました")
                    self.set_actions_state(True)
                    self.progress_reset()

                self.after(0, _on_error)
                return

            except Exception as e:
                def _on_error():
                    messagebox.showerror("エラー", f"圧縮中に予期しないエラーが発生しました:\n{e}")
                    self.status.set("圧縮に失敗しました")
                    self.set_actions_state(True)
                    self.progress_reset()

                self.after(0, _on_error)
                return

            def _on_success():
                self.progress_done()

                if report_lines:
                    msg = "圧縮結果:\n\n" + "\n".join(report_lines)
                else:
                    msg = "圧縮が完了しました。"

                if target_mb is not None:
                    msg = f"[目標サイズ: {target_mb:.2f}MB]\n\n" + msg

                messagebox.showinfo("完了", msg)
                self.status.set(f"圧縮を完了しました:{processed}件")
                self.set_actions_state(True)

                if self.open_after.get() and last_out_path is not None:
                    open_folder(last_out_path)

            self.after(0, _on_success)

        threading.Thread(target=worker, daemon=True).start()

    # ======================= 変換タブ（Word / Excel） =======================
    def convert_tab(self):
        frame = self.tab_convert

        # ← メニューに戻る
        top_bar = ttk.Frame(frame)
        top_bar.pack(fill="x", padx=10, pady=(10, 0))
        ttk.Button(top_bar, text="← メニューに戻る", command=self.show_menu).pack(side="right")

        ttk.Label(
            frame,
            text="PDFを Word / Excel に変換します（テキスト＋表）。複数PDFをまとめて変換できます。"
        ).pack(anchor="w", padx=10, pady=10)

        # 対象PDF
        targets_frame = ttk.Frame(frame)
        targets_frame.pack(fill="x", padx=10, pady=(0, 10))

        ttk.Label(targets_frame, text="変換するPDF（複数選択可）:").pack(side="left")

        self.convert_files: list[Path] = []
        self.convert_label = tk.StringVar(value="（未選択）")

        btn_choose = ttk.Button(
            targets_frame,
            text="PDFを選択",
            command=self.choose_convert_pdfs,
        )
        btn_choose.pack(side="left")
        self.action_buttons.append(btn_choose)

        btn_clear = ttk.Button(
            targets_frame,
            text="クリア",
            command=self.clear_convert_pdfs,
        )
        btn_clear.pack(side="left", padx=5)
        self.action_buttons.append(btn_clear)

        ttk.Label(
            targets_frame,
            textvariable=self.convert_label,
        ).pack(side="left", padx=10)

        if DND_AVAILABLE:
            frame.drop_target_register(DND_FILES)
            frame.dnd_bind("<<Drop>>", self.on_drop_convert)

        # 出力形式
        type_frame = ttk.LabelFrame(frame, text="出力形式")
        type_frame.pack(fill="x", padx=10, pady=(0, 10))

        self.convert_to_word = tk.BooleanVar(value=True)
        self.convert_to_excel = tk.BooleanVar(value=True)

        ttk.Checkbutton(
            type_frame,
            text="Word（.docx）に変換",
            variable=self.convert_to_word,
        ).pack(anchor="w", padx=5, pady=2)

        ttk.Checkbutton(
            type_frame,
            text="Excel（.xlsx）に変換（表は罫線付きで出力）",
            variable=self.convert_to_excel,
        ).pack(anchor="w", padx=5, pady=2)

        ttk.Label(
            frame,
            text=(
                "※ 表と認識された部分は、Excel上でセルごとに分割し、すべてのセルに罫線（枠線）を付けて出力します。\n"
                "※ Excel出力では、文字数に応じて列幅を自動調整し、セル内で折り返して全文が見えるように整えます。\n"
                "※ 画像しか含まれていないPDFは、文字や表を正しく抽出できない場合があります。"
            ),
            wraplength=800,
        ).pack(anchor="w", padx=10, pady=(0, 10))

        btn_run_convert = ttk.Button(
            frame,
            text="変換を実行",
            command=self.run_convert,
        )
        btn_run_convert.pack(pady=10)
        self.action_buttons.append(btn_run_convert)

    def _update_convert_label(self):
        if not self.convert_files:
            self.convert_label.set("（未選択）")
        elif len(self.convert_files) == 1:
            self.convert_label.set(self.convert_files[0].name)
        else:
            self.convert_label.set(
                f"{self.convert_files[0].name} ほか {len(self.convert_files) - 1}件"
            )

    def clear_convert_pdfs(self):
        self.convert_files = []
        self._update_convert_label()
        self.update_pdf_info(None)
        self.status.set("変換対象をクリアしました。")

    def on_drop_convert(self, event):
        pdf_paths = self._iter_dnd_pdf_paths(event)
        if not pdf_paths:
            return

        added = 0
        for path in pdf_paths:
            if path not in self.convert_files:
                self.convert_files.append(path)
                added += 1

        self._update_convert_label()
        self.status.set(f"D&Dで変換対象: {len(self.convert_files)} 件のPDFを選択しました。")

        if self.convert_files:
            self.update_pdf_info(self.convert_files[0])

    def choose_convert_pdfs(self):
        paths = filedialog.askopenfilenames(
            title="変換するPDFを選択",
            filetypes=[("PDFファイル", "*.pdf")],
        )
        if not paths:
            return

        self.convert_files = [Path(p) for p in paths]
        self._update_convert_label()
        self.status.set(f"変換対象: {len(self.convert_files)} 件のPDFを選択しました。")

        if self.convert_files:
            self.update_pdf_info(self.convert_files[0])

    @staticmethod
    def _convert_pdf_to_word(src: Path, out_path: Path) -> None:
        """
        pdfplumber でテキストと表を抽出し、Word に出力。
        （シンプル実装：ページごとに段落＋表を順に追加）
        """
        doc = Document()

        with pdfplumber.open(str(src)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    for line in text.splitlines():
                        doc.add_paragraph(line)

                tables = page.extract_tables()
                for tbl in tables or []:
                    if not tbl:
                        continue
                    rows = tbl
                    cols = max(len(r) for r in rows if r) if rows else 0
                    if cols == 0:
                        continue

                    w_table = doc.add_table(rows=len(rows), cols=cols)
                    for r_idx, row in enumerate(rows):
                        if not row:
                            continue
                        for c_idx, cell_val in enumerate(row):
                            w_table.cell(r_idx, c_idx).text = (cell_val or "").strip()
                    doc.add_paragraph("")  # 表のあとに空行

                if page_index < len(pdf.pages):
                    doc.add_page_break()

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))

    @staticmethod
    def _convert_pdf_to_excel(src: Path, out_path: Path) -> None:
        """
        pdfplumber で PDF からテキストと表を抽出し、Excel に出力する。

        ・1枚目のシート「テキスト」：
            - 各ページのテキストをそのまま縦に書き出す
            - 罫線や列幅の自動調整は行わない
        ・2枚目以降のシート：
            - 表として認識されたものを 1 表 = 1 シートで出力
            - すべてのセルに罫線（枠線）を付ける
            - 文字数に応じて列幅を自動調整し、折り返し表示
        """
        wb = Workbook()

        # --- 1枚目のシート：全テキスト用 ---
        ws_text = wb.active
        ws_text.title = "テキスト"

        text_row = 1  # テキストを書き込む行位置

        # 罫線スタイル（表シート用）
        thin = Side(style="thin")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        table_sheet_index = 0  # 何枚目の表か（シート名用）

        with pdfplumber.open(str(src)) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                # ---------- テキスト（1枚目のシート） ----------
                text = page.extract_text() or ""

                # ページ見出し
                ws_text.cell(row=text_row, column=1,
                             value="このタブでは読み込んだ文字を全て出力しています。表としての出力は他タブをご確認ください。")
                text_row += 1
                ws_text.cell(row=text_row, column=1, value=f" {page_idx}ページ目")
                text_row += 1

                if text.strip():
                    for line in text.splitlines():
                        ws_text.cell(row=text_row, column=1, value=line)
                        text_row += 1

                # ページ間を1行空ける
                text_row += 1

                # ---------- 表（2枚目以降のシート） ----------
                tables = page.extract_tables()
                if not tables:
                    continue

                for t in tables:
                    if not t:
                        continue

                    table_sheet_index += 1
                    ws_tbl = wb.create_sheet(title=f"表{table_sheet_index}")

                    # 列ごとの最大文字数（このシート内でのみ計算）
                    max_len_per_col: dict[int, int] = {}

                    for r_idx, row in enumerate(t, start=1):
                        if row is None:
                            continue
                        for c_idx, cell_val in enumerate(row, start=1):
                            val = (cell_val or "").strip()
                            cell = ws_tbl.cell(row=r_idx, column=c_idx, value=val)
                            cell.border = border
                            cell.alignment = Alignment(wrap_text=True, vertical="top")

                            if val:
                                # 1セル内の最長行の文字数を幅計算に使う
                                max_line_len = max(len(line) for line in val.splitlines())
                                prev = max_len_per_col.get(c_idx, 0)
                                if max_line_len > prev:
                                    max_len_per_col[c_idx] = max_line_len

                    # 列幅を自動調整（この表シートだけ）
                    for col_idx, length in max_len_per_col.items():
                        col_letter = get_column_letter(col_idx)
                        # ざっくり：1文字あたり 1.2 幅、最低 8、最大 80 に制限
                        width = max(8, min(80, length * 1.2))
                        ws_tbl.column_dimensions[col_letter].width = width

        # 保存
        out_path.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(out_path))

    def run_convert(self):
        if not getattr(self, "convert_files", None):
            self.convert_files = []

        if not self.convert_files:
            messagebox.showwarning("警告", "変換するPDFを選んでください。")
            return

        to_word = self.convert_to_word.get()
        to_excel = self.convert_to_excel.get()

        if not (to_word or to_excel):
            messagebox.showwarning("警告", "少なくとも1つは出力形式（Word / Excel）を選択してください。")
            return

        dir_str = self.output_dir_var.get().strip()

        tasks: list[tuple[Path, Optional[Path], Optional[Path]]] = []
        # （src, word_path, excel_path）

        for src in self.convert_files:
            src = Path(src)

            if dir_str:
                base_dir = Path(dir_str)
            else:
                base_dir = src.parent

            word_path = base_dir / f"{src.stem}.docx" if to_word else None
            excel_path = base_dir / f"{src.stem}.xlsx" if to_excel else None

            # 上書き確認（一括設定に従う）
            if word_path and not self.confirm_overwrite(word_path):
                word_path = None
            if excel_path and not self.confirm_overwrite(excel_path):
                excel_path = None

            if not word_path and not excel_path:
                # このファイルは全部スキップ
                continue

            tasks.append((src, word_path, excel_path))

        if not tasks:
            self.status.set("変換をキャンセルしました（すべてスキップ）")
            return

        self.set_actions_state(False)
        self.progress_reset()
        self.status.set("変換を開始しました...")

        total = len(tasks)

        def worker():
            last_out: Optional[Path] = None

            try:
                for i, (src, w_path, x_path) in enumerate(tasks, start=1):
                    # Word
                    if w_path:
                        try:
                            self._convert_pdf_to_word(src, w_path)
                            last_out = w_path
                        except Exception as e:
                            msg = f"{src.name} の Word 変換に失敗しました:\n{e}"
                            self.after(0, lambda m=msg: messagebox.showerror("エラー", m))

                    # Excel（表は罫線付き＆セルサイズ調整）
                    if x_path:
                        try:
                            self._convert_pdf_to_excel(src, x_path)
                            last_out = x_path
                        except Exception as e:
                            msg = f"{src.name} の Excel 変換に失敗しました:\n{e}"
                            self.after(0, lambda m=msg: messagebox.showerror("エラー", m))

                    percent = i / total * 100

                    def _update(p=percent, name=src.name, idx=i):
                        self.progress_set(p)
                        self.status.set(f"変換中...（{idx}/{total}）{name}")

                    self.after(0, _update)

            except Exception as e:
                def _on_error():
                    messagebox.showerror("エラー", f"変換中に予期しないエラーが発生しました:\n{e}")
                    self.status.set("変換に失敗しました")
                    self.set_actions_state(True)
                    self.progress_reset()

                self.after(0, _on_error)
                return

            def _on_success():
                self.progress_done()
                messagebox.showinfo("完了", f"変換が完了しました。（{len(tasks)} 件）")
                self.status.set("変換を完了しました。")
                self.set_actions_state(True)

                if self.open_after.get() and last_out is not None:
                    open_folder(last_out)

            self.after(0, _on_success)

        threading.Thread(target=worker, daemon=True).start()

    # ======================= パスワードタブ =======================
    def password_tab(self):
        frame = self.tab_password

        # ← メニューに戻る
        top_bar = ttk.Frame(frame)
        top_bar.pack(fill="x", padx=10, pady=(10, 0))
        ttk.Button(top_bar, text="← メニューに戻る", command=self.show_menu).pack(side="right")

        ttk.Label(
            frame,
            text=(
                "選択したPDFにパスワードを設定／解除します（新規ファイルとして出力）。\n"
                "・パスワードを設定：編集するにはパスワードが必要なファイルを出力します\n"
                "・パスワードを解除：既にパスワード付きのPDFからパスワードを外します"
            )
        ).pack(anchor="w", padx=10, pady=10)

        src_frame = ttk.Frame(frame)
        src_frame.pack(fill="x", padx=10, pady=(0, 10))

        self.lock_file: Optional[Path] = None
        self.lock_file_label = tk.StringVar(value="(未選択)")

        btn_choose = ttk.Button(
            src_frame,
            text="PDFを選択",
            command=self.choose_lock_pdf,
        )
        btn_choose.pack(side="left")
        self.action_buttons.append(btn_choose)

        btn_clear = ttk.Button(
            src_frame,
            text="クリア",
            command=self.clear_lock_pdf,
        )
        btn_clear.pack(side="left", padx=5)
        self.action_buttons.append(btn_clear)

        ttk.Label(
            src_frame,
            textvariable=self.lock_file_label,
        ).pack(side="left", padx=10)

        if DND_AVAILABLE:
            frame.drop_target_register(DND_FILES)
            frame.dnd_bind("<<Drop>>", self.on_drop_lock)

        mode_frame = ttk.LabelFrame(frame, text="処理モード")
        mode_frame.pack(fill="x", padx=10, pady=(0, 10))

        self.lock_mode = tk.StringVar(value="set")

        ttk.Radiobutton(
            mode_frame,
            text="パスワードを設定",
            value="set",
            variable=self.lock_mode,
            command=self.update_lock_mode_ui,
        ).pack(side="left", padx=5, pady=5)

        ttk.Radiobutton(
            mode_frame,
            text="パスワードを解除",
            value="clear",
            variable=self.lock_mode,
            command=self.update_lock_mode_ui,
        ).pack(side="left", padx=5, pady=5)

        pwd_frame = ttk.LabelFrame(frame, text="パスワード")
        pwd_frame.pack(fill="x", padx=10, pady=(0, 10))

        self.lock_pwd_var = tk.StringVar()
        self.lock_pwd_confirm_var = tk.StringVar()

        row1 = ttk.Frame(pwd_frame)
        row1.pack(fill="x", pady=3)
        ttk.Label(row1, text="パスワード:").pack(side="left", padx=5)
        self.lock_pwd_entry = ttk.Entry(row1, textvariable=self.lock_pwd_var, show="*")
        self.lock_pwd_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))

        row2 = ttk.Frame(pwd_frame)
        row2.pack(fill="x", pady=3)
        ttk.Label(row2, text="確認用:").pack(side="left", padx=5)
        self.lock_pwd_confirm_entry = ttk.Entry(
            row2, textvariable=self.lock_pwd_confirm_var, show="*"
        )
        self.lock_pwd_confirm_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))

        print_frame = ttk.Frame(frame)
        print_frame.pack(fill="x", padx=10, pady=(0, 10))

        self.disable_print = tk.BooleanVar(value=False)

        self.chk_disable_print = ttk.Checkbutton(
            print_frame,
            text="印刷を禁止する（閲覧のみ許可）",
            variable=self.disable_print,
        )
        self.chk_disable_print.pack(anchor="w")

        out_frame = ttk.Frame(frame)
        out_frame.pack(fill="x", padx=10, pady=(0, 10))

        ttk.Label(out_frame, text="出力ファイル名:").pack(side="left")

        self.lock_output_var = tk.StringVar(value="")
        self.lock_output_entry = ttk.Entry(
            out_frame,
            textvariable=self.lock_output_var,
            width=50,
        )
        self.lock_output_entry.pack(side="left", padx=5)

        self.init_placeholder(
            self.lock_output_entry,
            "空欄:'元ファイル名'_ロック済み.pdf / _解除済み.pdf",
        )

        self.update_lock_output_placeholder()

        btn_run_lock = ttk.Button(frame, text="パスワード処理を実行", command=self.run_lock_action)
        btn_run_lock.pack(pady=10)
        self.action_buttons.append(btn_run_lock)

        self.update_lock_mode_ui()

    def update_lock_mode_ui(self):
        mode = self.lock_mode.get()

        if mode == "set":
            self.chk_disable_print.configure(state="normal")
            self.lock_pwd_confirm_entry.config(state="normal")
        else:
            self.chk_disable_print.configure(state="disabled")
            self.lock_pwd_confirm_entry.config(state="disabled")

        self.lock_pwd_var.set("")
        self.lock_pwd_confirm_var.set("")
        self.update_lock_output_placeholder()

    def clear_lock_pdf(self):
        self.lock_file = None
        self.lock_file_label.set("(未選択)")
        self.update_pdf_info(None)
        self.update_lock_output_placeholder()
        self.status.set("パスワード対象をクリアしました。")

    def on_drop_lock(self, event):
        pdf_paths = self._iter_dnd_pdf_paths(event)
        if not pdf_paths:
            return

        path = pdf_paths[0]

        self.lock_file = path
        self.lock_file_label.set(path.name)
        self.status.set(f"パスワード対象（D&D）: {path}")
        self.update_pdf_info(path)
        self.update_lock_output_placeholder()

    def choose_lock_pdf(self):
        path = filedialog.askopenfilename(
            title="パスワード設定／解除をするPDFを選択",
            filetypes=[("PDFファイル", "*.pdf")],
        )
        if not path:
            return

        p = Path(path)
        self.lock_file = p
        self.lock_file_label.set(p.name)
        self.status.set(f"パスワード対象: {p}")
        self.update_pdf_info(p)
        self.update_lock_output_placeholder()

    def run_lock_action(self):
        if not getattr(self, "lock_file", None):
            messagebox.showwarning("警告", "対象PDFを選択してください。")
            return

        src = self.lock_file
        mode = self.lock_mode.get()
        pwd = self.lock_pwd_var.get()
        pwd2 = self.lock_pwd_confirm_var.get()

        dir_str = self.output_dir_var.get().strip()
        if dir_str:
            out_dir = Path(dir_str)
        else:
            out_dir = src.parent

        try:
            out_dir.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            messagebox.showerror("エラー", f"出力フォルダの作成に失敗しました:\n{e}")
            return

        if mode == "set":
            if not pwd:
                messagebox.showwarning("警告", "パスワードを入力してください。")
                return
            if pwd != pwd2:
                messagebox.showwarning("警告", "確認用パスワードが一致しません。")
                return

            raw_name = self.get_entry_text(self.lock_output_entry)
            if not raw_name:
                raw_name = self.get_lock_default_name("set", src)

            if not raw_name.lower().endswith(".pdf"):
                raw_name += ".pdf"

            out_path = out_dir / raw_name

            if not self.confirm_overwrite(out_path):
                self.status.set("パスワード設定をキャンセルしました（既存ファイルあり）")
                return

            try:
                allow_print = not self.disable_print.get()
                self.set_pdf_password(src, out_path, pwd, allow_print)
            except Exception as e:
                messagebox.showerror("エラー", f"パスワード設定に失敗しました:\n{e}")
                self.status.set("パスワード設定に失敗しました")
                return

            messagebox.showinfo("完了", f"パスワード付きPDFを出力しました:\n{out_path}")
            self.status.set(f"パスワード設定を完了しました: {out_path}")
            if self.open_after.get():
                open_folder(out_path)

        else:
            if not pwd:
                messagebox.showwarning("警告", "解除には現在のパスワードが必要です。")
                return

            raw_name = self.get_entry_text(self.lock_output_entry)
            if not raw_name:
                raw_name = self.get_lock_default_name("remove", src)

            if not raw_name.lower().endswith(".pdf"):
                raw_name += ".pdf"

            out_path = out_dir / raw_name
            if not self.confirm_overwrite(out_path):
                self.status.set("パスワード解除をキャンセルしました（既存ファイルあり）")
                return

            try:
                self.remove_pdf_password(src, out_path, pwd)
            except ValueError as e:
                messagebox.showerror("エラー", str(e))
                self.status.set("パスワード解除に失敗しました（パスワード不一致）")
                return
            except Exception as e:
                messagebox.showerror("エラー", f"パスワード解除に失敗しました:\n{e}")
                self.status.set("パスワード解除に失敗しました")
                return

            messagebox.showinfo("完了", f"パスワードを解除したPDFを出力しました:\n{out_path}")
            self.status.set(f"パスワード解除を完了しました: {out_path}")
            if self.open_after.get():
                open_folder(out_path)
