"""
PDF Conversion Service

Provides functionality to convert PDFs to Word and Excel formats.
"""

from pathlib import Path
from typing import Optional, Callable, List, Tuple

import pdfplumber
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment
from openpyxl.utils import get_column_letter


class PDFConvertError(Exception):
    """Exception raised during PDF conversion operations."""
    pass


def _convert_pdf_to_word(src: Path, out_path: Path) -> None:
    """
    Convert PDF to Word document.
    
    Extracts text and tables from PDF and creates a Word document.
    
    Args:
        src: Source PDF file path
        out_path: Output Word file path
        
    Raises:
        PDFConvertError: If conversion fails
    """
    try:
        doc = Document()
        
        with pdfplumber.open(str(src)) as pdf:
            total_pages = len(pdf.pages)
            
            for page_index, page in enumerate(pdf.pages, start=1):
                # Extract and add text
                text = page.extract_text() or ""
                if text.strip():
                    for line in text.splitlines():
                        if line.strip():  # Skip empty lines
                            doc.add_paragraph(line)
                
                # Extract and add tables
                tables = page.extract_tables()
                if tables:
                    for tbl in tables:
                        if not tbl:
                            continue
                        
                        # Calculate table dimensions
                        rows = tbl
                        cols = max(len(r) for r in rows if r) if rows else 0
                        if cols == 0:
                            continue
                        
                        # Create Word table
                        w_table = doc.add_table(rows=len(rows), cols=cols)
                        
                        # Fill table cells
                        for r_idx, row in enumerate(rows):
                            if not row:
                                continue
                            for c_idx, cell_val in enumerate(row):
                                if c_idx < cols:
                                    cell_text = (cell_val or "").strip()
                                    w_table.cell(r_idx, c_idx).text = cell_text
                        
                        doc.add_paragraph("")  # Add spacing after table
                
                # Add page break (except after last page)
                if page_index < total_pages:
                    doc.add_page_break()
        
        # Save document
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        
    except Exception as e:
        raise PDFConvertError(
            f"Word変換中にエラーが発生しました: {str(e)}"
        ) from e


def _convert_pdf_to_excel(src: Path, out_path: Path) -> None:
    """
    Convert PDF to Excel spreadsheet.
    
    Creates separate sheets for text and tables.
    
    Args:
        src: Source PDF file path
        out_path: Output Excel file path
        
    Raises:
        PDFConvertError: If conversion fails
    """
    try:
        wb = Workbook()
        
        # Create text sheet
        ws_text = wb.active
        ws_text.title = "テキスト"
        text_row = 1
        
        # Define cell border style
        thin = Side(style="thin")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        
        table_sheet_index = 0
        
        # Add header to text sheet
        ws_text.cell(
            row=text_row,
            column=1,
            value="このシートでは読み込んだ文字を全て出力しています。"
                  "表としての出力は別シートをご確認ください。",
        )
        text_row += 2
        
        with pdfplumber.open(str(src)) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                # Extract text
                text = page.extract_text() or ""
                
                # Add page header
                ws_text.cell(row=text_row, column=1, value=f"{page_idx}ページ目")
                text_row += 1
                
                # Add text content
                if text.strip():
                    for line in text.splitlines():
                        ws_text.cell(row=text_row, column=1, value=line)
                        text_row += 1
                
                text_row += 1  # Add spacing between pages
                
                # Extract tables
                tables = page.extract_tables()
                if not tables:
                    continue
                
                for tbl in tables:
                    if not tbl:
                        continue
                    
                    # Create new sheet for this table
                    table_sheet_index += 1
                    ws_tbl = wb.create_sheet(title=f"表{table_sheet_index}")
                    
                    # Track column widths
                    max_len_per_col: dict[int, int] = {}
                    
                    # Fill table data
                    for r_idx, row in enumerate(tbl, start=1):
                        if row is None:
                            continue
                        
                        for c_idx, cell_val in enumerate(row, start=1):
                            val = (cell_val or "").strip()
                            
                            # Create and style cell
                            cell = ws_tbl.cell(row=r_idx, column=c_idx, value=val)
                            cell.border = border
                            cell.alignment = Alignment(
                                wrap_text=True,
                                vertical="top"
                            )
                            
                            # Track max line length for column width
                            if val:
                                max_line_len = max(
                                    len(line) for line in val.splitlines()
                                )
                                prev = max_len_per_col.get(c_idx, 0)
                                if max_line_len > prev:
                                    max_len_per_col[c_idx] = max_line_len
                    
                    # Set column widths
                    for col_idx, length in max_len_per_col.items():
                        col_letter = get_column_letter(col_idx)
                        # Calculate width (with min/max limits)
                        width = max(8, min(80, length * 1.2))
                        ws_tbl.column_dimensions[col_letter].width = width
        
        # Save workbook
        out_path.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(out_path))
        
    except Exception as e:
        raise PDFConvertError(
            f"Excel変換中にエラーが発生しました: {str(e)}"
        ) from e


def convert_pdfs(
    tasks: List[Tuple[Path, Optional[Path], Optional[Path]]],
    progress_cb: Optional[Callable[[float, str], None]] = None,
) -> int:
    """
    Convert multiple PDFs to Word and/or Excel.
    
    Args:
        tasks: List of (src_pdf, word_out_path, excel_out_path) tuples
               Set word_out_path or excel_out_path to None to skip that format
        progress_cb: Optional callback receiving (percent, message)
        
    Returns:
        Number of files successfully converted
        
    Raises:
        PDFConvertError: If conversion fails
        
    Examples:
        >>> # Convert to both Word and Excel
        >>> tasks = [(src, word_out, excel_out)]
        >>> count = convert_pdfs(tasks)
        >>> # Convert to Word only
        >>> tasks = [(src, word_out, None)]
        >>> count = convert_pdfs(tasks)
    """
    total = len(tasks)
    completed = 0
    
    for idx, (src, word_path, excel_path) in enumerate(tasks, start=1):
        try:
            # Convert to Word if requested
            if word_path:
                _convert_pdf_to_word(src, word_path)
            
            # Convert to Excel if requested
            if excel_path:
                _convert_pdf_to_excel(src, excel_path)
            
            completed += 1
            
            # Report progress
            if progress_cb:
                percent = idx / total * 100
                progress_cb(percent, f"変換完了...（{idx}/{total}）{src.name}")
                
        except Exception as e:
            # Log error but continue with other files
            if progress_cb:
                progress_cb(
                    (idx - 1) / total * 100,
                    f"エラー: {src.name}"
                )
            raise
    
    return completed
